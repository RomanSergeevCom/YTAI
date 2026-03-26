#!/usr/bin/env python3
"""
YTAI 0507 Pre-Edit: Parse filled Google Doc (.docx) back to pre_edit.json

Reads a .docx that was exported by export_to_gdoc.py and filled by the user.
Extracts the "Визуализация" annotations per segment and classifies them.

Usage:
    python parse_gdoc.py --input filled_template.docx
    python parse_gdoc.py --input filled_template.docx --output Setup/Pre-Edit/v1/
"""

import argparse
import json
import sys
import os
import re
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Annotation classification
# ---------------------------------------------------------------------------

def classify_annotation(text):
    """Classify a visual annotation by its prefix/content."""
    text = text.strip()
    if not text:
        return {'type': 'talking_head', 'content': ''}

    # Remove hint line if present
    hint_pattern = r'/схема:.*\|.*/оверлей:.*\|.*B-roll:.*'
    text = re.sub(hint_pattern, '', text).strip()
    if not text:
        return {'type': 'talking_head', 'content': ''}

    lower = text.lower()

    if lower.startswith('/схема:') or lower.startswith('/schema:'):
        content = text.split(':', 1)[1].strip()
        return {'type': 'schema', 'content': content}

    if lower.startswith('/оверлей:') or lower.startswith('/overlay:'):
        content = text.split(':', 1)[1].strip()
        # Detect overlay subtype
        subtype = 'full_overlay'
        for keyword, st in [('half', 'half_overlay'), ('1/2', 'half_overlay'),
                            ('3/5', 'three_fifths_overlay'), ('lower', 'lower_third'),
                            ('chapter', 'chapter_bar'), ('bar', 'chapter_bar')]:
            if keyword in lower:
                subtype = st
                break
        return {'type': 'overlay', 'subtype': subtype, 'content': content}

    if lower.startswith('/найди кадр') or lower.startswith('/find frame'):
        content = text.split(' ', 2)[-1].strip() if len(text.split(' ', 2)) > 2 else text
        return {'type': 'find_frame', 'content': content}

    if lower.startswith('/структурируй:') or lower.startswith('/structure:'):
        content = text.split(':', 1)[1].strip()
        return {'type': 'structure', 'content': content}

    if lower.startswith('b-roll:') or lower.startswith('broll:'):
        content = text.split(':', 1)[1].strip()
        return {'type': 'broll', 'content': content}

    # Check for URLs
    if re.match(r'https?://', text):
        return {'type': 'video_ref', 'content': text}

    # Default: treat as free-form note
    return {'type': 'note', 'content': text}


def parse_docx(input_path):
    """Parse a filled .docx and extract annotations per segment."""
    doc = Document(input_path)

    segments = []
    current_segment = None
    in_visual_section = False

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()

                # Detect segment info line (seg_XXX)
                seg_match = re.search(r'(seg_\d{3})', cell_text)
                if seg_match and '│' in cell_text:
                    # This is an info cell — start of a new segment
                    if current_segment and in_visual_section:
                        # Save previous segment
                        segments.append(current_segment)

                    seg_id = seg_match.group(1)

                    # Extract info from the line
                    parts = cell_text.split('│')
                    source_file = parts[1].strip() if len(parts) > 1 else ''
                    tc_range = parts[2].strip() if len(parts) > 2 else ''
                    speaker = parts[4].strip() if len(parts) > 4 else ''

                    # Extract transcript (next paragraph after info line)
                    paragraphs = cell.paragraphs
                    transcript_lines = []
                    for p in paragraphs[1:]:  # Skip first (info line)
                        t = p.text.strip()
                        if t and not t.startswith('B-roll:') and not t.startswith('Notes:'):
                            transcript_lines.append(t)

                    current_segment = {
                        'segment_id': seg_id,
                        'source_file': source_file.strip(),
                        'tc_range': tc_range.strip(),
                        'speaker': speaker.strip(),
                        'transcript': ' '.join(transcript_lines),
                        'visual_raw': '',
                        'visual': None,
                    }
                    in_visual_section = False
                    continue

                # Detect visualization section
                if 'Визуализация:' in cell_text:
                    in_visual_section = True
                    # Extract the content after "Визуализация:"
                    vis_lines = []
                    for p in cell.paragraphs:
                        t = p.text.strip()
                        if t and t != 'Визуализация:' and not t.startswith('/схема:') or (t.startswith('/') and ':' in t):
                            # Include /commands but skip the hint line
                            if '(пусто = talking head)' not in t:
                                vis_lines.append(t)

                    # Remove "Визуализация:" prefix
                    raw = '\n'.join(vis_lines).replace('Визуализация:', '').strip()
                    # Remove hint line
                    raw = re.sub(r'/схема: \.\.\. \| /оверлей: \.\.\. \|.*', '', raw).strip()

                    if current_segment:
                        current_segment['visual_raw'] = raw
                        current_segment['visual'] = classify_annotation(raw)

                    # Check for embedded images
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                                if current_segment:
                                    current_segment['visual'] = {
                                        'type': 'image',
                                        'content': 'Embedded image in document',
                                        'has_embedded_image': True,
                                    }

    # Don't forget the last segment
    if current_segment:
        segments.append(current_segment)

    return segments


def extract_images(input_path, output_dir):
    """Extract embedded images from .docx to output directory."""
    doc = Document(input_path)
    images = []

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            ext = rel.target_part.content_type.split('/')[-1]
            if ext == 'jpeg':
                ext = 'jpg'
            img_name = f"embedded_{len(images) + 1}.{ext}"
            img_path = Path(output_dir) / img_name
            img_path.parent.mkdir(parents=True, exist_ok=True)
            with open(img_path, 'wb') as f:
                f.write(image_data)
            images.append(str(img_path))

    return images


def main():
    parser = argparse.ArgumentParser(description='Parse filled Pre-Edit .docx to JSON')
    parser.add_argument('--input', '-i', required=True, help='Path to filled .docx')
    parser.add_argument('--output', '-o', help='Output directory (default: Pre-Edit/v{N}/ next to input)')
    parser.add_argument('--version', '-v', type=int, help='Version number (default: auto-detect)')
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Parse document
    segments = parse_docx(str(input_path))
    print(f"Parsed {len(segments)} segments from {input_path.name}")

    # Count annotation types
    type_counts = {}
    for s in segments:
        if s['visual']:
            t = s['visual']['type']
            type_counts[t] = type_counts.get(t, 0) + 1
    print(f"Annotations: {type_counts}")

    # Determine output directory
    if args.output:
        output_dir = Path(args.output).resolve()
    else:
        parent = input_path.parent
        # Auto-detect version
        version = args.version
        if not version:
            existing = [d for d in parent.iterdir() if d.is_dir() and d.name.startswith('v')]
            max_v = 0
            for d in existing:
                try:
                    v = int(d.name[1:])
                    if v > max_v:
                        max_v = v
                except ValueError:
                    pass
            version = max_v + 1

        output_dir = parent / f'v{version}'

    output_dir.mkdir(parents=True, exist_ok=True)
    visuals_dir = output_dir / 'visuals'
    visuals_dir.mkdir(exist_ok=True)

    # Extract embedded images
    images = extract_images(str(input_path), str(visuals_dir))
    if images:
        print(f"Extracted {len(images)} embedded images")

    # Match images to segments that have image annotations
    img_idx = 0
    for s in segments:
        if s['visual'] and s['visual'].get('has_embedded_image') and img_idx < len(images):
            s['visual']['image_path'] = images[img_idx]
            img_idx += 1

    # Build output JSON
    code = ''
    code_match = re.search(r'(YT[A-Z]{2,4}\d+)', input_path.name)
    if code_match:
        code = code_match.group(1)

    output_data = {
        'project_code': code,
        'source_docx': str(input_path),
        'parsed_at': datetime.now().isoformat(),
        'version': version if not args.output else 1,
        'segments': [
            {
                'segment_id': s['segment_id'],
                'source_file': s['source_file'],
                'tc_range': s['tc_range'],
                'speaker': s['speaker'],
                'transcript': s['transcript'],
                'visual_type': s['visual']['type'] if s['visual'] else 'talking_head',
                'visual_content': s['visual'].get('content', '') if s['visual'] else '',
                'visual_subtype': s['visual'].get('subtype', '') if s['visual'] else '',
                'visual_image_path': s['visual'].get('image_path', '') if s['visual'] else '',
            }
            for s in segments
        ],
        'stats': {
            'total_segments': len(segments),
            'annotations': type_counts,
            'embedded_images': len(images),
        }
    }

    # Write JSON
    json_name = f'{code}_pre_edit_v{output_data["version"]}.json' if code else 'pre_edit.json'
    json_path = output_dir / json_name
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)

    print(f"Written: {json_path}")
    print(f"Next: python generate_visuals.py --input \"{json_path}\"")


if __name__ == '__main__':
    main()
