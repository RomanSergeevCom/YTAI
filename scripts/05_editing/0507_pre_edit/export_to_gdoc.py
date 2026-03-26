#!/usr/bin/env python3
"""
YTAI 0507 Pre-Edit: Export brief to Google Doc-compatible .docx

Reads a pre_edit_export.json (from UXP Export or from Assembly _in.json)
and generates a .docx with:
  - Header: project name, sequence, date
  - Per-block colored headers
  - Per-segment rows: info + transcript + empty "Визуализация" section
  - The user opens this in Google Docs, fills column "Визуализация"

Usage:
    python export_to_gdoc.py --input YTCR01_Assembly_v4_in.json
    python export_to_gdoc.py --input Setup/Pre-Edit/YTCR01_pre_edit_export.json
    python export_to_gdoc.py --input YTCR01_Assembly_v4_in.json --output custom_path.docx
"""

import argparse
import json
import sys
import os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Color mapping (block color name → RGB)
# ---------------------------------------------------------------------------
BLOCK_COLORS = {
    'Green':   RGBColor(0x2E, 0x7D, 0x32),
    'Blue':    RGBColor(0x2A, 0x5A, 0x8A),
    'Cyan':    RGBColor(0x00, 0x80, 0x7E),
    'Yellow':  RGBColor(0x9E, 0x8A, 0x00),
    'Red':     RGBColor(0xA3, 0x28, 0x2E),
    'Magenta': RGBColor(0x8E, 0x1E, 0x8E),
    'Orange':  RGBColor(0xB8, 0x7A, 0x20),
    'Purple':  RGBColor(0x6A, 0x3D, 0x7D),
}

VISUAL_PLACEHOLDER = (
    "Визуализация: идея поиска | схема/инфографика | Gemini промпт | Claude анимация | (пусто = talking head)"
)


def tc_to_sec(tc):
    """Parse MM:SS.s timecode to seconds."""
    parts = tc.split(':')
    if len(parts) == 2:
        return int(parts[0]) * 60 + float(parts[1])
    return 0


def format_duration(sec):
    """Format seconds as M:SS."""
    m = int(sec) // 60
    s = int(sec) % 60
    return f"{m}:{s:02d}"


def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.find(qn('w:shd'))
    if shading_elm is None:
        from docx.oxml import OxmlElement
        shading_elm = OxmlElement('w:shd')
        shading.append(shading_elm)
    shading_elm.set(qn('w:fill'), hex_color)
    shading_elm.set(qn('w:val'), 'clear')


def find_clean_transcript(input_path):
    """Search for _clean_transcript.json near the input file."""
    input_dir = Path(input_path).parent
    # Search patterns: same dir, Assembly/ sibling, parent Setup/
    search_dirs = [
        input_dir,
        input_dir.parent / 'Assembly',
        input_dir.parent,
    ]
    for d in search_dirs:
        if not d.is_dir():
            continue
        for f in d.glob('*_clean_transcript.json'):
            return f
    return None


def load_segments(input_path):
    """Load segments from either _in.json brief or _pre_edit_export.json.
    Enriches with clean transcript if available."""
    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    segments = data.get('segments', [])
    if not segments:
        print(f"ERROR: No segments found in {input_path}", file=sys.stderr)
        sys.exit(1)

    # Load clean transcript if available
    clean_lookup = {}
    clean_path = find_clean_transcript(input_path)
    if clean_path:
        with open(clean_path, 'r', encoding='utf-8') as f:
            clean_data = json.load(f)
        for cs in clean_data.get('segments', []):
            key = cs.get('segment_id', '')
            if not key:
                # Match by source_file + tc_in
                key = cs.get('source_file', '') + '|' + cs.get('tc_in', '')
            clean_lookup[key] = cs
        print(f"Clean transcript: {clean_path.name} ({len(clean_lookup)} segments)")

    # Normalize to common format
    result = []
    for s in segments:
        # Skip unused segments
        use = s.get('use', s.get('use', 'TRUE'))
        if str(use).upper() == 'FALSE':
            continue

        seg_id = s.get('segment_id', s.get('id', ''))
        source_file = s.get('source_file', s.get('sourceFile', ''))
        tc_in = s.get('tc_in', '')

        # Look up clean version — validate that clean_text matches the segment's transcript
        clean = clean_lookup.get(seg_id) or clean_lookup.get(source_file + '|' + tc_in) or {}
        raw_tx = s.get('transcript', '')
        clean_tx = clean.get('clean_text', '')
        # Only use clean_text if it shares first 20 chars with the raw transcript
        # (prevents stale clean_transcript from overriding corrected briefs)
        if clean_tx and raw_tx:
            raw_start = raw_tx.lower().replace(' ', '')[:20]
            clean_start = clean_tx.lower().replace(' ', '')[:20]
            if raw_start[:10] not in clean_start and clean_start[:10] not in raw_start:
                print(f"  WARNING: clean_transcript mismatch for {seg_id}, using brief transcript", file=sys.stderr)
                clean = {}  # discard entire stale clean entry
                clean_tx = ''

        seg = {
            'segment_id': seg_id,
            'source_file': source_file,
            'tc_in': tc_in,
            'tc_out': s.get('tc_out', ''),
            'block': s.get('block', 0),
            'block_name': s.get('block_name', s.get('blockName', '')),
            'speaker': s.get('speaker', ''),
            'color': s.get('color', 'Green'),
            'transcript': clean_tx or raw_tx,
            'summary': clean.get('summary', ''),
            'visual_hints': clean.get('visual_hints', []),
            'broll_note': s.get('broll_note', s.get('brollNote', '')),
            'notes': s.get('notes', ''),
            'visual': s.get('visual', ''),
            'visual_type': s.get('visual_type', ''),
            'visual_file': s.get('visual_file', ''),
            'marker_comments': s.get('marker_comments', []),
        }

        # Compute duration
        if s.get('duration'):
            seg['duration'] = s['duration']
        elif seg['tc_in'] and seg['tc_out']:
            seg['duration'] = round(tc_to_sec(seg['tc_out']) - tc_to_sec(seg['tc_in']), 1)
        else:
            seg['duration'] = 0

        result.append(seg)

    # Extract project info
    project = data.get('project', {})
    project_name = project.get('project_name', '') or data.get('project', '') or data.get('projectName', '')
    sequence = data.get('sequence', '')

    return result, project_name, sequence


def set_cell_text(cell, text, size=9, color=None, bold=False, italic=False):
    """Set cell text with formatting. Clears existing content."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def set_col_widths(table, widths_cm):
    """Set fixed column widths on a table (works in Pages/Google Docs/Word)."""
    from docx.oxml import OxmlElement

    # 1. Disable autofit
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 2. Set table layout to fixed (prevents auto-redistribution)
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    # 3. Set total table width
    total_twips = sum(int(w * 567) for w in widths_cm)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total_twips))
    tblW.set(qn('w:type'), 'dxa')

    # 4. Set tblGrid (authoritative column width definitions in OOXML)
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    else:
        for child in list(tblGrid):
            tblGrid.remove(child)
    for width in widths_cm:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(width * 567)))
        tblGrid.append(gridCol)

    # 5. Set each cell width explicitly
    for i, width in enumerate(widths_cm):
        twips = str(int(width * 567))
        for row in table.rows:
            if i < len(row.cells):
                cell = row.cells[i]
                tc = cell._element.get_or_add_tcPr()
                tcW = tc.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tc.append(tcW)
                tcW.set(qn('w:w'), twips)
                tcW.set(qn('w:type'), 'dxa')
                # Vertical align top
                vAlign = tc.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = OxmlElement('w:vAlign')
                    tc.append(vAlign)
                vAlign.set(qn('w:val'), 'top')


# Column definitions: name, width_cm
# 4 columns: compact # + info stack, wide transcript, wide visualization
COLUMNS = [
    ('#',             1.0),
    ('info',          2.5),
    ('транскрипт',   15.0),
    ('визуализация',  9.2),
]

COL_NAMES = [c[0] for c in COLUMNS]
COL_WIDTHS = [c[1] for c in COLUMNS]


def build_docx(segments, project_name, sequence, output_path):
    """Generate .docx document with columnar pre-edit table."""
    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(8)

    # --- Landscape orientation for wide table ---
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.orientation = 1  # Landscape
    # Swap width/height for landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    # --- Title ---
    title = doc.add_heading(level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f'{project_name} — Pre-Edit')
    run.font.color.rgb = RGBColor(0x26, 0x32, 0x38)
    run.font.size = Pt(14)

    sub = doc.add_paragraph()
    run = sub.add_run(f'{sequence} | {datetime.now().strftime("%Y-%m-%d %H:%M")} | {len(segments)} segments')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # --- Instructions (compact) ---
    instr = doc.add_paragraph()
    run = instr.add_run(
        'Визуализация: идея поиска | схема/инфографика | Gemini промпт | Claude анимация | (пусто = talking head)'
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x60, 0x7D, 0x8B)
    run.italic = True

    # --- Single continuous table ---
    INFO_COLOR = RGBColor(0x78, 0x90, 0x9C)
    current_block = None
    seg_num = 0

    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Column headers row
    hdr_row = table.rows[0]
    for ci, col_name in enumerate(COL_NAMES):
        cell = hdr_row.cells[ci]
        set_cell_text(cell, col_name, size=7, bold=True, color=RGBColor(0x55, 0x55, 0x55))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, 'EEEEEE')

    set_col_widths(table, COL_WIDTHS)

    for seg in segments:
        block_key = (seg['block'], seg['block_name'])
        seg_num += 1

        # New block → colored chapter separator row (merged)
        if block_key != current_block:
            current_block = block_key
            color = seg.get('color', 'Green')
            rgb = BLOCK_COLORS.get(color, BLOCK_COLORS['Green'])
            hex_color = str(rgb)

            block_row = table.add_row()
            block_row.cells[0].merge(block_row.cells[-1])
            merged_cell = block_row.cells[0]
            merged_cell.text = ''
            p = merged_cell.paragraphs[0]
            block_label = f'{seg["block"]}. {seg["block_name"]}' if seg['block_name'] else f'Block {seg["block"]}'
            run = p.add_run(block_label)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(merged_cell, hex_color)

        # --- Segment row ---
        row = table.add_row()
        cells = row.cells

        # Subtle alternating row tint
        if seg_num % 2 == 0:
            for c in cells:
                set_cell_bg(c, 'FAFAFA')

        # Col 0: # (segment number, bold, centered)
        cells[0].text = ''
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(seg_num))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = INFO_COLOR

        # Col 1: info (stacked: TC, duration, speaker)
        cells[1].text = ''
        p = cells[1].paragraphs[0]
        run = p.add_run(f'{seg["tc_in"]}→\n{seg["tc_out"]}')
        run.font.size = Pt(7)
        run.font.color.rgb = INFO_COLOR
        p2 = cells[1].add_paragraph()
        run = p2.add_run(format_duration(seg.get('duration', 0)))
        run.font.size = Pt(7)
        run.font.color.rgb = INFO_COLOR
        run.bold = True
        speaker = seg.get('speaker', '')
        if speaker:
            short_speaker = speaker.replace('Speaker ', 'Sp.')
            p3 = cells[1].add_paragraph()
            run = p3.add_run(short_speaker)
            run.font.size = Pt(7)
            run.font.color.rgb = INFO_COLOR

        # Col 2: transcript (clean text only — no B-roll, no visual hints)
        cells[2].text = ''
        transcript = seg.get('transcript', '') or ''
        summary = seg.get('summary', '') or ''

        if summary:
            p_sum = cells[2].paragraphs[0]
            p_sum.paragraph_format.space_before = Pt(0)
            p_sum.paragraph_format.space_after = Pt(4)
            run = p_sum.add_run(summary)
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x60, 0x7D, 0x8B)
            run.italic = True

        paragraphs = transcript.split('\n\n') if '\n\n' in transcript else [transcript]
        for pi, para_text in enumerate(paragraphs):
            para_text = para_text.strip()
            if not para_text:
                continue
            if summary or pi > 0:
                p_tx = cells[2].add_paragraph()
            else:
                p_tx = cells[2].paragraphs[0]
            p_tx.paragraph_format.space_before = Pt(2)
            p_tx.paragraph_format.space_after = Pt(2)
            run = p_tx.add_run(para_text)
            run.font.size = Pt(9)

        # Col 3: визуализация (B-roll, visual hints, ideas, Gemini prompts)
        vis_cell = cells[3]
        vis_cell.text = ''

        existing_visual = seg.get('visual', '') or ''
        visual_type = seg.get('visual_type', '') or ''
        visual_hints = seg.get('visual_hints', []) or []
        marker_comments = seg.get('marker_comments', []) or []
        broll = seg.get('broll_note', '')
        notes = seg.get('notes', '')

        has_content = False

        # B-roll note (moved from transcript column)
        if broll:
            p_br = vis_cell.paragraphs[0] if not has_content else vis_cell.add_paragraph()
            p_br.paragraph_format.space_before = Pt(1)
            p_br.paragraph_format.space_after = Pt(2)
            run = p_br.add_run(broll)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
            run.bold = True
            has_content = True

        # Visual hints from clean transcript
        if visual_hints:
            for hint in visual_hints:
                p_h = vis_cell.add_paragraph() if has_content else vis_cell.paragraphs[0]
                p_h.paragraph_format.space_before = Pt(1)
                p_h.paragraph_format.space_after = Pt(1)
                run = p_h.add_run(hint)
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)
                run.italic = True
                has_content = True

        # Existing visual (from previous Pre-Edit round)
        if existing_visual or visual_type:
            p_v = vis_cell.add_paragraph() if has_content else vis_cell.paragraphs[0]
            prefix = f'[{visual_type}] ' if visual_type else ''
            run = p_v.add_run(f'{prefix}{existing_visual}')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            has_content = True

        # Marker /comments (orange bold)
        if marker_comments:
            for mc in marker_comments:
                p_mc = vis_cell.add_paragraph() if has_content else vis_cell.paragraphs[0]
                run = p_mc.add_run(mc)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xB8, 0x7A, 0x20)
                run.bold = True
                has_content = True

        # Cell background tint
        if marker_comments:
            set_cell_bg(vis_cell, 'FFF8E1')
        elif has_content:
            set_cell_bg(vis_cell, 'F0FFF0')
        else:
            set_cell_bg(vis_cell, 'F5F9FF')

    # --- Save ---
    doc.save(output_path)
    print(f"Written: {output_path}")
    print(f"Segments: {len(segments)}")


def main():
    parser = argparse.ArgumentParser(description='Export Assembly brief to Google Doc-compatible .docx')
    parser.add_argument('--input', '-i', required=True, help='Path to _in.json brief or _pre_edit_export.json')
    parser.add_argument('--output', '-o', help='Output .docx path (default: same dir as input, _pre_edit_template.docx)')
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    segments, project_name, sequence = load_segments(str(input_path))

    if args.output:
        output_path = Path(args.output).resolve()
    else:
        # Auto-resolve output path
        parent = input_path.parent
        # If we're in Assembly/ folder, go up to Setup/ and into Pre-Edit/
        if parent.name == 'Assembly':
            pre_edit_dir = parent.parent / 'Pre-Edit'
        else:
            pre_edit_dir = parent
        pre_edit_dir.mkdir(parents=True, exist_ok=True)

        # Extract project code
        code = project_name or input_path.stem.split('_')[0]
        output_path = pre_edit_dir / f'{code}_pre_edit_template.docx'

    build_docx(segments, project_name, sequence, str(output_path))


if __name__ == '__main__':
    main()
